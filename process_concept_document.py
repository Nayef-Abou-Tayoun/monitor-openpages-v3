#!/usr/bin/env python3
"""
Process Concept Document with Watson Orchestrate Agent
Similar workflow to find_process.py and upload_summaries_from_cos.py:
1. Extract text from DOCX
2. Send to Watson Orchestrate agent
3. Format response as DOCX
4. Save outputs (JSON + DOCX)
5. Optionally upload to OpenPages
"""

import os
import sys
import json
import base64
import io
import re
import asyncio
import nest_asyncio
from datetime import datetime
from typing import Optional, Dict, Any
import requests
import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
import pytz

# Apply nest_asyncio to allow asyncio.run() in existing event loops
nest_asyncio.apply()

# Load environment variables
load_dotenv()

# Watson Orchestrate Configuration
WXO_API_KEY = os.getenv("WXO_API_KEY")
WXO_INSTANCE_ID = os.getenv("WXO_INSTANCE_ID")
WXO_AGENT_ID = os.getenv("WXO_AGENT_ID")

# OpenPages Configuration (optional)
OPENPAGES_SERVER = os.getenv("OPENPAGES_SERVER")
OPENPAGES_USERNAME = os.getenv("OPENPAGES_USERNAME")
OPENPAGES_PASSWORD = os.getenv("OPENPAGES_PASSWORD")
PROCESS_ID = os.getenv("PROCESS_ID")
PROCESS_NAME = os.getenv("PROCESS_NAME", "Concept Document Review")


class ConceptDocumentProcessor:
    """Process concept documents with Watson Orchestrate"""
    
    def __init__(self):
        self.wxo_api_key = WXO_API_KEY
        self.wxo_instance_id = WXO_INSTANCE_ID
        self.wxo_agent_id = WXO_AGENT_ID
        self.token = None
        self.token_expiry = None
        
    def log(self, message: str):
        """Print timestamped log message"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"[{timestamp}] {message}")
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract ALL text from DOCX or TXT file, including tables and nested content"""
        try:
            # Check file extension
            if file_path.lower().endswith('.txt'):
                # Read text file directly
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted = f.read()
                self.log(f"✅ Extracted {len(extracted)} characters from {file_path}")
                return extracted
            elif file_path.lower().endswith('.docx'):
                # Extract from DOCX - get EVERYTHING
                doc = Document(file_path)
                text = []
                
                # Method 1: Extract all paragraphs (gets most content)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text.append(para.text)
                
                # Method 2: Extract all tables (including nested tables)
                def extract_table_text(table):
                    table_text = ["\n[TABLE]"]
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            # Get all text from cell including nested tables
                            cell_content = []
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    cell_content.append(paragraph.text.strip())
                            # Check for nested tables in cell
                            for nested_table in cell.tables:
                                cell_content.append(extract_table_text(nested_table))
                            if cell_content:
                                row_text.append(" ".join(cell_content))
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    table_text.append("[/TABLE]\n")
                    return "\n".join(table_text)
                
                for table in doc.tables:
                    text.append(extract_table_text(table))
                
                extracted = '\n'.join(text)
                self.log(f"✅ Extracted {len(extracted)} characters from {file_path} (complete extraction)")
                return extracted
            else:
                self.log(f"❌ Unsupported file type: {file_path}")
                return ""
        except Exception as e:
            self.log(f"❌ Error extracting text: {e}")
            import traceback
            traceback.print_exc()
            return ""
    
    def get_iam_token(self) -> Optional[str]:
        """Get IAM token for Watson Orchestrate authentication"""
        try:
            # Use Watson Orchestrate specific token endpoint
            url = "https://iam.platform.saas.ibm.com/siusermgr/api/1.0/apikeys/token"
            
            response = requests.post(url, json={"apikey": self.wxo_api_key}, timeout=30)
            
            if response.status_code == 200:
                token_data = response.json()
                self.token = token_data.get("token")
                # Token expires in 600 seconds (10 minutes) by default
                self.token_expiry = datetime.now().timestamp() + token_data.get("expires_in", 600)
                self.log(f"✅ Token obtained (expires in {token_data.get('expires_in', 600)}s)")
                return self.token
            else:
                self.log(f"❌ Failed to get IAM token: HTTP {response.status_code}")
                self.log(f"   Response: {response.text[:200]}")
                return None
        except Exception as e:
            self.log(f"❌ Error getting IAM token: {e}")
            return None
    
    def is_token_expired(self) -> bool:
        """Check if token is expired"""
        if not self.token or not self.token_expiry:
            return True
        return datetime.now().timestamp() >= self.token_expiry - 300  # Refresh 5 min before expiry
    
    def get_bearer_token(self) -> Optional[str]:
        """Get valid bearer token, refreshing if needed"""
        if self.is_token_expired():
            return self.get_iam_token()
        return self.token
    
    def trigger_wxo_agent(self, document_text: str, filename: str, doc_id: str = "DOC-001", 
                          process_id: str = "PROC-001") -> Optional[Dict[str, Any]]:
        """Trigger Watson Orchestrate agent with document content"""
        try:
            # Get bearer token
            token = self.get_bearer_token()
            if not token:
                self.log("❌ Failed to get Watson Orchestrate token")
                return None
            
            # Prepare the prompt for the agent - send ALL content (no truncation)
            prompt = f"""{document_text}"""
            
            # Watson Orchestrate chat completions API
            url = f"https://api.dl.watson-orchestrate.ibm.com/instances/{self.wxo_instance_id}/v1/orchestrate/{self.wxo_agent_id}/chat/completions"
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "stream": False
            }
            
            self.log("🤖 Triggering Watson Orchestrate agent...")
            self.log(f"   Agent ID: {self.wxo_agent_id}")
            
            response = requests.post(url, headers=headers, json=payload, timeout=120)
            
            if response.status_code == 200:
                data = response.json()
                summary = data.get("choices", [{}])[0].get("message", {}).get("content", "No response returned.")
                
                self.log("✅ Agent response received successfully!")
                self.log(f"   Response length: {len(summary)} characters")
                
                # Use EST timezone for timestamp
                est = pytz.timezone('US/Eastern')
                est_time = datetime.now(est)
                
                # Get agent name from response metadata if available
                agent_name = data.get("model", f"Agent_{self.wxo_agent_id[:8]}")
                
                result = {
                    "summary": summary,
                    "document_name": filename,
                    "document_id": doc_id,
                    "process_id": process_id,
                    "process_name": PROCESS_NAME,
                    "timestamp": est_time.strftime("%d/%m/%Y %H:%M:%S"),
                    "agent_name": agent_name,
                    "agent_id": self.wxo_agent_id,
                    "metadata": {
                        "text_length": len(document_text),
                        "word_count": len(document_text.split()),
                        "extraction_method": "python-docx",
                        "ai_model": f"Watson Orchestrate - {agent_name}",
                        "generated_by": "Real Watson Orchestrate API"
                    }
                }
                return result
            else:
                self.log(f"❌ Watson Orchestrate agent failed: HTTP {response.status_code}")
                self.log(f"   Response: {response.text[:500]}")
                return None
        
        except Exception as e:
            self.log(f"❌ Exception calling Watson Orchestrate agent: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with markdown-style formatting (bold, italic)"""
        parts = []
        current = ""
        i = 0
        
        while i < len(text):
            if i < len(text) - 1 and text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                # Find closing **
                end = text.find('**', i + 2)
                if end != -1:
                    parts.append(('bold', text[i+2:end]))
                    i = end + 2
                    continue
            current += text[i]
            i += 1
        
        if current:
            parts.append(('normal', current))
        
        # Add runs to paragraph
        for style, content in parts:
            run = paragraph.add_run(content)
            if style == 'bold':
                run.font.bold = True
    
    def _parse_markdown_table(self, lines: list, start_idx: int) -> tuple:
        """Parse a markdown table and return (table_data, end_idx)"""
        table_rows = []
        i = start_idx
        
        while i < len(lines):
            line = lines[i].strip()
            if not line.startswith('|'):
                break
            # Skip separator lines (---|---|---)
            if '---' in line:
                i += 1
                continue
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
            if cells:
                table_rows.append(cells)
            i += 1
        
        return table_rows, i
    
    def _format_cell_text(self, cell, text: str, is_header: bool = False):
        """Format cell text, handling bold markers properly"""
        from docx.shared import Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Clear existing content
        cell.text = ''
        paragraph = cell.paragraphs[0]
        
        # Set cell to not break words (no wrap)
        tcPr = cell._element.get_or_add_tcPr()
        noWrap = OxmlElement('w:noWrap')
        tcPr.append(noWrap)
        
        # Process text to handle **bold** markers
        parts = []
        current = ""
        i = 0
        
        while i < len(text):
            if i < len(text) - 1 and text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                # Find closing **
                end = text.find('**', i + 2)
                if end != -1:
                    parts.append(('bold', text[i+2:end]))
                    i = end + 2
                    continue
            current += text[i]
            i += 1
        
        if current:
            parts.append(('normal', current))
        
        # Add runs with proper formatting
        for style, content in parts:
            run = paragraph.add_run(content)
            if style == 'bold' or is_header:
                run.font.bold = True
            run.font.size = Pt(9)
    
    def _add_docx_table(self, doc, table_data: list):
        """Add a formatted table to the document with proper formatting"""
        from docx.shared import Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        if not table_data or len(table_data) < 2:
            return
        
        # Create table
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Prevent table from breaking across pages
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Add keep together property
        keepTogether = OxmlElement('w:tblpPr')
        tblPr.append(keepTogether)
        
        # Set equal column widths for all tables
        num_cols = len(table_data[0])
        # Equal widths for all columns
        widths = [Inches(7.0 / num_cols)] * num_cols
        
        # Set column widths and prevent page breaks within table rows
        for row_idx, row in enumerate(table.rows):
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width
            
            # Prevent page break within this row
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)
            
            # For header row (first row), mark it as a repeating header and keep with next
            if row_idx == 0:
                # Mark as table header (repeats on each page if table spans pages)
                tblHeader = OxmlElement('w:tblHeader')
                tblHeader.set(qn('w:val'), '1')
                trPr.append(tblHeader)
                
                # Keep header with next row (prevents orphaned header)
                keepNext = OxmlElement('w:keepNext')
                trPr.append(keepNext)
        
        # Fill table with formatted content
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                is_header = (row_idx == 0)
                self._format_cell_text(cell, cell_text, is_header)
    
    def format_summary_as_docx(self, wxo_result: Dict[str, Any], filename: str) -> Optional[bytes]:
        """Format Watson Orchestrate summary as a DOCX document with proper table parsing"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('WATSON ORCHESTRATE RISK ASSESSMENT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document Information Section
            doc.add_heading('Document Information', level=1)
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_data = [
                ('Document Name:', filename),
                ('Document ID:', str(wxo_result.get('document_id', 'N/A'))),
                ('Process ID:', str(wxo_result.get('process_id', 'N/A'))),
                ('Process Name:', wxo_result.get('process_name', 'N/A')),
                ('Generated:', wxo_result.get('timestamp', 'N/A')),
                ('Agent:', wxo_result.get('agent_name', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = value
                # Bold the labels
                info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()  # Add spacing
            
            summary = wxo_result.get('summary', 'No summary available')
            
            # Import re for regex matching
            import re
            
            # Parse the summary - the Executive Summary will be handled by the ## heading parser
            lines = summary.split('\n')
            
            # Initialize section counter (not used anymore since we extract from agent's output)
            section_number = 1
            
            # Parse and format the summary
            current_paragraph = None
            skip_section = True  # Start by skipping until we reach a section we want
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    current_paragraph = None
                    i += 1
                    continue
                
                # Check if this is the start of a markdown table
                if line.startswith('|'):
                    # Skip table if we're in a section to skip
                    if skip_section:
                        # Find end of table
                        j = i + 1
                        while j < len(lines) and lines[j].strip().startswith('|'):
                            j += 1
                        i = j
                        continue
                    table_data, end_idx = self._parse_markdown_table(lines, i)
                    if table_data:
                        self._add_docx_table(doc, table_data)
                        doc.add_paragraph()  # Add spacing after table
                    i = end_idx
                    current_paragraph = None
                    continue
                
                # Skip separator lines
                if line.startswith('---'):
                    i += 1
                    continue
                
                # Headers with ##
                if line.startswith('##'):
                    header_text = line.replace('##', '').strip()
                    # Skip "Relevant Risk Categories", "Risk Assessment", "Query Statistics" sections
                    if ('relevant risk categories' in header_text.lower() or
                        'risk assessment' in header_text.lower() or
                        'query statistics' in header_text.lower()):
                        skip_section = True
                        current_paragraph = None
                    elif 'executive summary' in header_text.lower():
                        skip_section = False  # Don't skip - we want this section
                        # Extract number from header like "1.0 Executive Summary"
                        import re
                        match = re.match(r'^(\d+\.0)\s+', header_text)
                        if match:
                            section_num = match.group(1)
                        else:
                            section_num = '1.0'
                        doc.add_heading(f'{section_num} Executive Summary', level=1)
                        current_paragraph = None
                    elif 'openpages risk analysis' in header_text.lower():
                        skip_section = False  # Reset skip flag - we want this section
                        # Use the section number from the agent's output (e.g., "2.0" or "3.0")
                        # Extract number from header like "2.0 OpenPages Risk Analysis" or "3.0 OpenPages Risk Analysis"
                        import re
                        match = re.match(r'^(\d+\.0)\s+', header_text)
                        if match:
                            section_num = match.group(1)
                        else:
                            section_num = f'{section_number}.0'
                            section_number += 1
                        doc.add_heading(f'{section_num} OpenPages Risk Analysis', level=1)
                        current_paragraph = None
                    elif 'openpages control analysis' in header_text.lower():
                        skip_section = False
                        # Extract number from header
                        import re
                        match = re.match(r'^(\d+\.0)\s+', header_text)
                        if match:
                            section_num = match.group(1)
                        else:
                            section_num = f'{section_number}.0'
                            section_number += 1
                        doc.add_heading(f'{section_num} OpenPages Control Analysis', level=1)
                        current_paragraph = None
                    elif 'summary and recommendations' in header_text.lower():
                        skip_section = False
                        # Extract number from header
                        import re
                        match = re.match(r'^(\d+\.0)\s+', header_text)
                        if match:
                            section_num = match.group(1)
                        else:
                            section_num = f'{section_number}.0'
                            section_number += 1
                        doc.add_heading(f'{section_num} Summary and Recommendations', level=1)
                        current_paragraph = None
                    elif 'query statistics' in header_text.lower():
                        # Skip the entire Query Statistics section
                        skip_section = True
                        current_paragraph = None
                    else:
                        skip_section = False  # Reset skip flag for new sections
                        doc.add_heading(header_text, level=2)
                        current_paragraph = None
                # Headers with ### (level 3 headings)
                elif line.startswith('###'):
                    header_text = line.replace('###', '').strip()
                    header_lower = header_text.lower()
                    
                    # Skip unwanted sections
                    if ('relevant risk categories' in header_lower or
                        'business risk types' in header_lower or
                        'groups impacted' in header_lower or
                        'categories impacted' in header_lower):
                        skip_section = True
                        current_paragraph = None
                    else:
                        skip_section = False
                        doc.add_heading(header_text, level=3)
                        current_paragraph = None
                # Headers with **text**
                elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
                    header_text = line.replace('**', '').strip()
                    header_lower = header_text.lower()
                    
                    # Skip unwanted sections including ANY query statistics
                    if ('relevant risk categories' in header_lower or
                        'new initiative risk assessment' in header_lower or
                        'display only' in header_lower or
                        'business risk types' in header_lower or
                        'groups impacted' in header_lower or
                        'categories impacted' in header_lower or
                        'project summary' in header_lower or  # Skip standalone Project Summary heading
                        'query statistics' in header_lower or
                        'queried' in header_lower or  # Catches any heading with "queried"
                        'identified' in header_lower):  # Catches any heading with "identified"
                        skip_section = True
                        current_paragraph = None
                    else:
                        skip_section = False  # Reset skip flag for new sections
                        doc.add_heading(header_text, level=2)
                        current_paragraph = None
                # Subheadings with *text:* (like *Project Summary:* or *Brief project summary*:)
                elif line.startswith('*') and (':*' in line or '*:' in line) and line.count('*') >= 2:
                    # Extract text between * and :*
                    header_text = line.replace('*', '').replace(':', '').strip()
                    
                    # Skip Project Summary variations - already added as a separate section
                    # Also skip "Groups Impacted" or "Categories Impacted"
                    if ('project summary' in header_text.lower() or
                        'groups impacted' in header_text.lower() or
                        'categories impacted' in header_text.lower()):
                        current_paragraph = None
                    else:
                        # Other subheadings use normal heading style
                        doc.add_heading(header_text, level=3)
                        current_paragraph = None
                # Numbered items (but check if it's actually a heading first)
                elif len(line) > 2 and line[0:2].replace('.', '').replace(')', '').isdigit():
                    if not skip_section:
                        # Check if this looks like a heading (contains bold markers or common heading words)
                        if '**' in line or any(word in line.lower() for word in ['privacy', 'monitoring', 'security', 'legal', 'governance', 'compliance', 'technology', 'risk']):
                            # Treat as a subheading - keep the number but remove ** markers
                            heading_text = line.replace('**', '')
                            doc.add_heading(heading_text, level=3)
                            current_paragraph = None
                        else:
                            # Regular numbered list item
                            p = doc.add_paragraph(style='List Number')
                            self._add_formatted_text(p, line)
                            current_paragraph = None
                # Bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    # Check if this bullet point contains query statistics
                    line_lower = line.lower()
                    skip_bullet = False
                    
                    query_keywords = [
                        'total risks queried',
                        'relevant risks identified',
                        'total controls queried',
                        'relevant controls identified',
                        'query statistics',
                        'risks queried',
                        'controls queried',
                        'risks identified',
                        'controls identified',
                        'queried:',
                        'identified:',
                        'out of',  # Catches "20 out of 137 total controls"
                        'total controls returned',
                        'total risks returned',
                    ]
                    
                    for keyword in query_keywords:
                        if keyword in line_lower:
                            skip_bullet = True
                            break
                    
                    if not skip_bullet and not skip_section:
                        p = doc.add_paragraph(style='List Bullet')
                        self._add_formatted_text(p, line[2:])
                        current_paragraph = None
                # Regular text
                else:
                    # Skip query statistics lines - remove formatting first then check
                    skip_line = False
                    
                    # Remove all formatting markers before checking
                    line_clean = line.replace('*', '').replace(':', '').strip().lower()
                    
                    # Regex patterns that match lines containing these word combinations
                    query_patterns = [
                        r'(?i).*risks?.*queried.*',
                        r'(?i).*controls?.*queried.*',
                        r'(?i).*risks?.*identified.*',
                        r'(?i).*controls?.*identified.*',
                        r'(?i).*query.*statistics.*',
                        r'(?i).*items?.*marked.*',
                    ]
                    
                    # Check if cleaned line matches any query pattern
                    for pattern in query_patterns:
                        if re.match(pattern, line_clean):
                            skip_line = True
                            break
                    
                    if skip_line:
                        current_paragraph = None
                    elif not skip_section:
                        if current_paragraph is None:
                            current_paragraph = doc.add_paragraph()
                            self._add_formatted_text(current_paragraph, line)
                        else:
                            current_paragraph.add_run(' ')
                            self._add_formatted_text(current_paragraph, line)
                
                i += 1
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            self.log("✅ DOCX document formatted successfully")
            return docx_bytes.getvalue()
        
        except Exception as e:
            self.log(f"❌ Error formatting DOCX: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def save_outputs(self, wxo_result: Dict[str, Any], docx_content: Optional[bytes], 
                     base_filename: str) -> tuple:
        """Save JSON and DOCX outputs"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save JSON
        json_filename = f"{base_filename}_{timestamp}.json"
        try:
            with open(json_filename, 'w', encoding='utf-8') as f:
                json.dump(wxo_result, f, indent=2, ensure_ascii=False)
            self.log(f"✅ Saved JSON output: {json_filename}")
        except Exception as e:
            self.log(f"❌ Error saving JSON: {e}")
            json_filename = None
        
        # Save DOCX
        docx_filename = None
        if docx_content:
            docx_filename = f"{base_filename}_{timestamp}.docx"
            try:
                with open(docx_filename, 'wb') as f:
                    f.write(docx_content)
                self.log(f"✅ Saved DOCX output: {docx_filename}")
            except Exception as e:
                self.log(f"❌ Error saving DOCX: {e}")
                docx_filename = None
        
        return json_filename, docx_filename
    
    async def upload_to_openpages(self, file_content: bytes, filename: str,
                                   description: str, process_resource_id: str = "31619",
                                   folder_id: str = "31628") -> Optional[str]:
        """Upload a document to OpenPages using httpx async client"""
        try:
            if not all([OPENPAGES_SERVER, OPENPAGES_USERNAME, OPENPAGES_PASSWORD]):
                self.log("⚠ OpenPages credentials not configured, skipping upload")
                return None
            
            self.log(f"📤 Uploading to OpenPages...")
            self.log(f"   Folder ID: {folder_id}")
            self.log(f"   Process ID: {process_resource_id}")
            
            # Create summary filename with timestamp to ensure uniqueness
            timestamp_suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
            name_without_ext = filename.rsplit('.', 1)[0]
            summary_filename = f"NIRA Report - {name_without_ext} - {timestamp_suffix}"
            
            # Encode file content
            file_content_b64 = base64.b64encode(file_content).decode('utf-8')
            
            # Remove /openpages from base_url if present
            base = OPENPAGES_SERVER.replace('/openpages', '').rstrip('/')
            create_url = f"{base}/grc/api/contents"
            
            create_payload = {
                "contentDefinition": {
                    "attribute": {
                        "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    },
                    "children": file_content_b64
                },
                "fileTypeDefinition": {
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                },
                "fields": {
                    "field": []
                },
                "typeDefinitionId": "4",  # SOXDocument type
                "parentFolderId": folder_id,
                "name": summary_filename,
                "description": description
            }
            
            # Use httpx async client
            async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                response = await http_client.post(
                    create_url,
                    json=create_payload,
                    auth=(OPENPAGES_USERNAME, OPENPAGES_PASSWORD),
                    timeout=60.0
                )
                
                if response.status_code in [200, 201]:
                    doc_data = response.json()
                    doc_id = doc_data.get('id')
                    self.log(f"   ✅ Document created (ID: {doc_id})")
                    
                    # Associate document with process
                    assoc_url = f"{base}/grc/api/contents/{process_resource_id}/associations/children"
                    assoc_payload = [{"id": doc_id}]
                    
                    self.log(f"   🔗 Associating with process...")
                    assoc_response = await http_client.post(
                        assoc_url,
                        json=assoc_payload,
                        auth=(OPENPAGES_USERNAME, OPENPAGES_PASSWORD),
                        timeout=30.0
                    )
                    
                    if assoc_response.status_code in [200, 201, 204]:
                        self.log(f"   ✅ Successfully associated with process")
                        return doc_id
                    else:
                        self.log(f"   ⚠ Association failed: HTTP {assoc_response.status_code}")
                        return doc_id
                else:
                    self.log(f"   ❌ Upload failed: HTTP {response.status_code}")
                    self.log(f"      {response.text[:200]}")
                    return None
                    
        except Exception as e:
            self.log(f"❌ Error uploading to OpenPages: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def extract_risk_count(self, summary_text: str) -> int:
        """Extract the number of relevant risks from the summary text"""
        import re
        # Look for patterns like "Relevant Risks Identified: 10" or "*Relevant Risks Identified: 10*"
        patterns = [
            r'Relevant Risks Identified:\s*(\d+)',
            r'\*Relevant Risks Identified:\s*(\d+)\*',
            r'Total Risks.*?Identified:\s*(\d+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, summary_text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        
        # If no pattern found, return 0
        return 0
    
    async def process_document(self, docx_path: str, doc_id: str = "DOC-001",
                        process_id: str = None, num_runs: int = 5) -> Optional[Dict[str, Any]]:
        """Main processing workflow - runs agent multiple times and selects best result"""
        print(f"\n{'='*70}")
        print(f"PROCESSING CONCEPT DOCUMENT WITH WATSON ORCHESTRATE")
        print(f"{'='*70}")
        print(f"Document: {docx_path}")
        print(f"Agent ID: {self.wxo_agent_id}")
        print(f"Number of runs: {num_runs}")
        print(f"{'='*70}\n")
        
        # Use environment variable if process_id not provided
        if not process_id:
            process_id = PROCESS_ID or "PROC-001"
        
        # Step 1: Extract text from file
        self.log("📄 Step 1: Extracting text from document...")
        document_text = self.extract_text_from_file(docx_path)
        if not document_text:
            self.log("❌ Failed to extract text from document")
            return None
        
        # Step 2: Run agent multiple times and collect results
        self.log(f"🤖 Step 2: Running Watson Orchestrate agent {num_runs} times...")
        filename = os.path.basename(docx_path)
        
        results = []
        for run_num in range(1, num_runs + 1):
            self.log(f"   Run {run_num}/{num_runs}...")
            wxo_result = self.trigger_wxo_agent(document_text, filename, doc_id, process_id)
            
            if wxo_result:
                risk_count = self.extract_risk_count(wxo_result.get('summary', ''))
                results.append({
                    'result': wxo_result,
                    'risk_count': risk_count,
                    'run_number': run_num
                })
                self.log(f"   ✅ Run {run_num} complete - {risk_count} risks identified")
            else:
                self.log(f"   ❌ Run {run_num} failed")
        
        if not results:
            self.log("❌ All runs failed to get response from Watson Orchestrate")
            return None
        
        # Step 3: Select result with maximum risks
        best_result = max(results, key=lambda x: x['risk_count'])
        self.log(f"\n🏆 Best result: Run {best_result['run_number']} with {best_result['risk_count']} risks")
        
        # Display summary of all runs
        self.log("\n📊 Summary of all runs:")
        for r in sorted(results, key=lambda x: x['risk_count'], reverse=True):
            self.log(f"   Run {r['run_number']}: {r['risk_count']} risks")
        
        wxo_result = best_result['result']
        self.log("\n✅ Selected best agent response")
        
        # Step 4: Format as DOCX
        self.log("📝 Step 3: Formatting response as DOCX...")
        docx_content = self.format_summary_as_docx(wxo_result, filename)
        
        # Step 5: Save outputs
        self.log("💾 Step 4: Saving outputs...")
        base_filename = filename.replace('.docx', '').replace(' ', '_')
        json_file, docx_file = self.save_outputs(wxo_result, docx_content, base_filename)
        
        # Step 6: Upload to OpenPages (if configured)
        uploaded_doc_id = None
        if docx_content and OPENPAGES_SERVER and OPENPAGES_USERNAME and OPENPAGES_PASSWORD:
            self.log("📤 Step 5: Uploading to OpenPages...")
            description = f"NIRA Report generated by Watson Orchestrate for {filename}"
            
            # Run async upload
            try:
                uploaded_doc_id = await self.upload_to_openpages(
                    docx_content,
                    filename,
                    description,
                    process_resource_id=process_id if process_id and process_id != "PROC-001" else "31619",
                    folder_id="31628"
                )
                if uploaded_doc_id:
                    self.log(f"✅ Successfully uploaded to OpenPages (Doc ID: {uploaded_doc_id})")
                else:
                    self.log("⚠ Upload to OpenPages failed")
            except Exception as e:
                self.log(f"❌ Error during upload: {str(e)}")
        else:
            self.log("⚠ Skipping OpenPages upload (not configured or no DOCX content)")
        
        print(f"\n{'='*70}")
        print(f"✅ PROCESSING COMPLETE")
        print(f"{'='*70}")
        print(f"🏆 Best Result: Run {best_result['run_number']} with {best_result['risk_count']} risks")
        if json_file:
            print(f"📄 JSON Output: {json_file}")
        if docx_file:
            print(f"📄 DOCX Output: {docx_file}")
        if uploaded_doc_id:
            print(f"📤 OpenPages Doc ID: {uploaded_doc_id}")
        print(f"{'='*70}\n")
        
        return wxo_result


async def main():
    """Main entry point"""
    # Check for document path argument
    if len(sys.argv) < 2:
        print("❌ Error: Document path not provided")
        print("\nUsage:")
        print("  python process_concept_document.py <path_to_document>")
        print("\nExample:")
        print("  python process_concept_document.py 'Example Concept document.docx'")
        print("  python process_concept_document.py 'Example_Concept_document.txt'")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    # Check if file exists
    if not os.path.exists(docx_path):
        print(f"❌ Error: File not found: {docx_path}")
        sys.exit(1)
    
    # Validate required environment variables
    if not all([WXO_API_KEY, WXO_INSTANCE_ID, WXO_AGENT_ID]):
        print("❌ Error: Watson Orchestrate credentials not configured")
        print("   Required: WXO_API_KEY, WXO_INSTANCE_ID, WXO_AGENT_ID")
        sys.exit(1)
    
    try:
        processor = ConceptDocumentProcessor()
        result = await processor.process_document(docx_path)
        
        if result:
            print("\n✅ Success! Document processed and outputs saved.")
            sys.exit(0)
        else:
            print("\n❌ Processing failed. Check logs above for details.")
            sys.exit(1)
    
    except KeyboardInterrupt:
        print("\n\n⏹ Processing stopped by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
    main()

# Made with Bob